from __future__ import annotations

import io
import logging
from typing import List, Tuple

from ...parser.models import (
    ParsedCell,
    ParsedTable,
    Provenance,
    TableRef,
    TextChunk,
    TextStyle,
)

logger = logging.getLogger(__name__)


class DocxExtractor:
    """Extract text and tables from Word documents using python-docx."""

    def extract(
        self,
        data: bytes,
        filename: str,
        provenance: Provenance,
    ) -> Tuple[str, List[ParsedTable]]:
        from docx import Document

        doc = Document(io.BytesIO(data))
        text_parts: List[str] = []
        tables: List[ParsedTable] = []

        # Paragraphs
        for para in doc.paragraphs:
            text = para.text.strip()
            if text:
                text_parts.append(text)

        # Tables
        for tbl_idx, tbl in enumerate(doc.tables):
            chain = provenance.table_chain.copy()
            chain.append(TableRef(table_index=tbl_idx, row=None, col=None))
            table_prov = Provenance(
                page_id=provenance.page_id,
                page_title=provenance.page_title,
                page_date=provenance.page_date,
                table_chain=chain,
                attachment_id=provenance.attachment_id,
                attachment_type=provenance.attachment_type,
            )

            headers: List[str] = []
            cells_matrix: List[List[ParsedCell]] = []

            for row_idx, row in enumerate(tbl.rows):
                row_cells: List[ParsedCell] = []
                for col_idx, cell in enumerate(row.cells):
                    text = cell.text.strip()
                    cell_prov = Provenance(
                        page_id=provenance.page_id,
                        page_title=provenance.page_title,
                        page_date=provenance.page_date,
                        table_chain=chain.copy(),
                        row=row_idx,
                        col=col_idx,
                        attachment_id=provenance.attachment_id,
                        attachment_type=provenance.attachment_type,
                    )
                    chunks = (
                        [TextChunk(content=text, style=TextStyle.NORMAL, provenance=cell_prov)]
                        if text else []
                    )
                    row_cells.append(ParsedCell(
                        row=row_idx, col=col_idx,
                        text_chunks=chunks,
                        provenance=cell_prov,
                    ))

                    # Flatten table cell text into the main text as well
                    if text:
                        text_parts.append(text)

                if row_idx == 0:
                    headers = [
                        c.text_chunks[0].content if c.text_chunks else ""
                        for c in row_cells
                    ]
                cells_matrix.append(row_cells)

            tables.append(ParsedTable(
                table_index=tbl_idx,
                headers=headers,
                cells=cells_matrix,
                provenance=table_prov,
                raw_html=f"[DOCX table {tbl_idx}]",
            ))

        return "\n".join(text_parts), tables
